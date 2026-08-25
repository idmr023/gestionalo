import os
import shutil
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma

from docx import Document as DocxDocument
from pypdf import PdfReader

DATA_ANSWER_DIR = 'api/data_answer'
CHROMA_PERSIST_DIR = 'api/chroma_db'
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY", "YOUR_GOOGLE_API_KEY")


def load_docx(file_path):
    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return [Document(page_content=text, metadata={"source": file_path})]


def load_pdf(file_path):
    reader = PdfReader(file_path)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": file_path, "page": i + 1}
            ))
    return docs


def process_documents():
    documents = []
    for root, _, files in os.walk(DATA_ANSWER_DIR):
        for file in files:
            file_path = os.path.join(root, file)
            print(f"Loading {file_path}...")
            if file.lower().endswith('.pdf'):
                documents.extend(load_pdf(file_path))
            elif file.lower().endswith('.docx'):
                documents.extend(load_docx(file_path))
            else:
                print(f"Skipping unsupported file: {file_path}")

    print(f"Loaded {len(documents)} raw documents.")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    texts = text_splitter.split_documents(documents)
    print(f"Split into {len(texts)} text chunks.")

    print("Creating embeddings and storing in ChromaDB...")
    embeddings = GoogleGenerativeAIEmbeddings(
        model="gemini-embedding-2-preview",
        google_api_key=GOOGLE_API_KEY,
    )

    if os.path.exists(CHROMA_PERSIST_DIR):
        shutil.rmtree(CHROMA_PERSIST_DIR)
        print(f"Removed existing ChromaDB at {CHROMA_PERSIST_DIR}")

    vectordb = Chroma.from_documents(
        documents=texts,
        embedding=embeddings,
        persist_directory=CHROMA_PERSIST_DIR
    )
    vectordb.persist()
    print(f"ChromaDB created and persisted at {CHROMA_PERSIST_DIR}")


if __name__ == '__main__':
    process_documents()